#!/usr/bin/env python3
"""Generate DOCX test fixtures for the OpenBrain DOCX ingest test harness.

Usage:
    python scripts/test_fixtures/generate_docx_fixtures.py

Requires: python-docx (prod dependency)
Output: scripts/test_fixtures/docx/*.docx
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent / "docx"

# Unique phrase embedded in simple_text.docx — used for retrieval smoke test.
# Do not change without updating smoke_checks.py and test_docx_url_ingest_eval.py.
SIMPLE_TEXT_UNIQUE_PHRASE = "xyloquartz-retrieval-fixture-beta"

# Known phrases embedded per fixture — used by eval harness.
# Keys must match fixture filenames (without .docx).
FIXTURE_KNOWN_PHRASES: dict[str, list[str]] = {
    "simple_text": [
        SIMPLE_TEXT_UNIQUE_PHRASE,
        "The quick brown fox jumps over the lazy dog",
    ],
    "multi_paragraph": [
        "Multi-paragraph fixture paragraph one begins here",
        "Multi-paragraph fixture paragraph eleven ends here",
    ],
    "special_chars": [
        "café naïve résumé",
        "Straße München",
    ],
}


def _require_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("ERROR: python-docx is required to generate fixtures.")
        print("Install with: pip install python-docx==1.2.0")
        sys.exit(1)


def generate_simple_text(out_dir: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("OpenBrain DOCX Test Fixture: simple_text")
    doc.add_paragraph(f"Unique retrieval phrase: {SIMPLE_TEXT_UNIQUE_PHRASE}")
    doc.add_paragraph(
        "The quick brown fox jumps over the lazy dog. "
        "This is a single-paragraph DOCX with clean text. "
        "It serves as the baseline extraction case for the OpenBrain test harness."
    )
    path = out_dir / "simple_text.docx"
    doc.save(str(path))
    return path


def generate_multi_paragraph(out_dir: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Multi-paragraph fixture paragraph one begins here. "
                       "This is the opening paragraph of the multi-paragraph DOCX fixture.")
    for i in range(2, 11):
        doc.add_paragraph(
            f"Multi-paragraph fixture paragraph {i}. "
            f"Each paragraph must be present in extraction output to verify "
            f"the extractor joins all paragraphs correctly."
        )
    doc.add_paragraph(
        "Multi-paragraph fixture paragraph eleven ends here. "
        "If this text is retrievable, all paragraphs were correctly extracted."
    )
    path = out_dir / "multi_paragraph.docx"
    doc.save(str(path))
    return path


def generate_empty(out_dir: Path) -> Path:
    from docx import Document

    doc = Document()
    # No paragraphs added — document body is empty.
    path = out_dir / "empty.docx"
    doc.save(str(path))
    return path


def generate_large(out_dir: Path) -> Path:
    from docx import Document

    # Default OPENBRAIN_TEXT_INGEST_MAX_WORDS is 6000 — generate ~7000 words.
    word = "infrastructure"
    sentence = " ".join([word] * 14) + "."  # 14 words + period = 15 tokens
    paragraph_text = " ".join([sentence] * 10)  # ~150 words per paragraph

    doc = Document()
    doc.add_paragraph("OpenBrain DOCX Test Fixture: large")
    # 47 paragraphs × ~150 words = ~7050 words
    for _ in range(47):
        doc.add_paragraph(paragraph_text)
    path = out_dir / "large.docx"
    doc.save(str(path))
    return path


def generate_special_chars(out_dir: Path) -> Path:
    from docx import Document

    # python-docx supports full Unicode — use real em-dashes, curly quotes, etc.
    doc = Document()
    doc.add_paragraph("OpenBrain DOCX Test Fixture: special_chars")
    doc.add_paragraph("French: café naïve résumé")
    doc.add_paragraph("German: Straße München über")
    doc.add_paragraph("Spanish: mañana camión año")
    doc.add_paragraph("Punctuation: em\u2014dash and \u201ccurly quotes\u201d")
    doc.add_paragraph("Math: 100° angle, π ≈ 3.14159")
    doc.add_paragraph("Currency: £50, ¥500, €99")
    path = out_dir / "special_chars.docx"
    doc.save(str(path))
    return path


GENERATORS = {
    "simple_text": generate_simple_text,
    "multi_paragraph": generate_multi_paragraph,
    "empty": generate_empty,
    "large": generate_large,
    "special_chars": generate_special_chars,
}


def main() -> int:
    _require_docx()
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    failed = 0
    for name, fn in GENERATORS.items():
        try:
            path = fn(OUT_DIR)
            size = os.path.getsize(path)
            print(f"PASS  {name}.docx ({size:,} bytes) -> {path}")
        except Exception as exc:
            print(f"FAIL  {name}.docx: {exc}")
            failed += 1

    if failed:
        print(f"\n{failed} fixture(s) failed to generate.")
    else:
        print(f"\nAll {len(GENERATORS)} fixtures generated in {OUT_DIR}")
    return failed


if __name__ == "__main__":
    raise SystemExit(main())
